from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ACCENT = RGBColor(26, 71, 122)
ACCENT_LIGHT = "DCE6F2"
TEXT = RGBColor(44, 44, 44)
MUTED = RGBColor(90, 90, 90)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = TEXT
            if row_index == 0:
                set_cell_shading(cell, ACCENT_LIGHT)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = ACCENT


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.color.rgb = ACCENT


def add_body(doc: Document, text: str, *, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    if bold_label:
        label = p.add_run(f"{bold_label}: ")
        label.bold = True
        label.font.color.rgb = TEXT
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F5F7FA")
    set_cell_margins(cell, top=110, start=140, bottom=110, end=140)
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(58, 58, 58)


def add_three_col_table(doc: Document, headers: list[str], rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.autofit = True
    hdr = table.rows[0].cells
    for idx, value in enumerate(headers):
        hdr[idx].text = value
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table)


def add_four_col_table(doc: Document, headers: list[str], rows: list[tuple[str, str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.autofit = True
    hdr = table.rows[0].cells
    for idx, value in enumerate(headers):
        hdr[idx].text = value
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    style_table(table)


def build_document(output_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = TEXT
    styles["Heading 1"].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.bold = True
    styles["Heading 2"].font.name = "Aptos Display"
    styles["Heading 2"].font.size = Pt(12.5)
    styles["Heading 2"].font.bold = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(60)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("AI Service Integration Handoff")
    run.font.name = "Aptos Display"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    srun = subtitle.add_run(
        "Backend API design and data-contract guide for connecting the main backend to the resume parsing and candidate matching AI services."
    )
    srun.font.size = Pt(11)
    srun.font.color.rgb = MUTED

    meta = doc.add_table(rows=3, cols=2)
    meta.cell(0, 0).text = "Project"
    meta.cell(0, 1).text = "AI-Resume-Screener"
    meta.cell(1, 0).text = "Prepared for"
    meta.cell(1, 1).text = "Main backend implementation handoff"
    meta.cell(2, 0).text = "Prepared on"
    meta.cell(2, 1).text = date.today().isoformat()
    style_table(meta)

    doc.add_paragraph().add_run()
    add_body(
        doc,
        "This document is based on the current `ai-service` codebase. It describes the AI-facing contracts that already exist, the backend endpoints that should wrap them, the model-field mappings the backend must perform, and the integration mismatches that should be normalized in one adapter layer."
    )

    doc.add_page_break()

    add_heading(doc, "1. What Exists Today", 1)
    add_body(doc, "The AI service currently exposes two primary public endpoints plus a health check.")
    add_three_col_table(
        doc,
        ["AI Endpoint", "Purpose", "Current Contract Notes"],
        [
            ("GET /health", "Health probe", "Returns `{status: ok, port}`."),
            ("POST /parse", "Resume parsing", "Multipart upload with `file`; returns a `ParsedCandidate` object."),
            ("POST /match/", "Candidate ranking", "JSON body with one job and many candidates; returns ranked candidates for that job."),
        ],
    )
    add_bullet(doc, "The AI service is a backend-to-backend dependency. Frontend clients should not call it directly.")
    add_bullet(doc, "The backend should own authentication, authorization, persistence, id translation, retries, and user-facing API shape.")
    add_bullet(doc, "The backend should treat the AI service as a pure compute/enrichment service.")

    add_heading(doc, "2. Recommended Ownership Split", 1)
    add_bullet(doc, "Backend owns uploads, user authorization, database writes, job and candidate lookups, and final API responses sent to the frontend.")
    add_bullet(doc, "AI service owns text extraction, resume parsing, embedding generation, similarity scoring, and natural-language ranking summaries.")
    add_bullet(doc, "Backend should never re-implement AI scoring logic in controller code; it should forward normalized payloads to `/match/`.")

    add_heading(doc, "3. AI Service Contracts", 1)
    add_heading(doc, "3.1 Parse Resume", 2)
    add_body(doc, "Multipart endpoint for resume parsing.", bold_label="AI route")
    add_code_block(doc, ["POST /parse", "Content-Type: multipart/form-data", "Body: file=<pdf|docx>"])
    add_body(doc, "Supported file rules and error behavior are below.", bold_label="Behavior")
    add_bullet(doc, "Accepts `.pdf` and `.docx` only.")
    add_bullet(doc, "Rejects files larger than `MAX_FILE_SIZE_MB`; current default is 5 MB.")
    add_bullet(doc, "Returns HTTP 400 for known app exceptions with `{success:false,message,error_code}`.")
    add_bullet(doc, "Returns HTTP 500 for unexpected failures.")
    add_body(doc, "The parser returns the fields listed below.", bold_label="Response shape")
    add_four_col_table(
        doc,
        ["Field", "Type", "Required", "Backend handling"],
        [
            ("full_name", "string", "No", "Store in `CandidateProfile.personalInfo.fullName`."),
            ("email", "string", "No", "Store in `CandidateProfile.personalInfo.email`."),
            ("phone", "string", "No", "Store in `CandidateProfile.personalInfo.phone`."),
            ("location", "string", "No", "Add to profile schema or keep in metadata if needed."),
            ("skills", "string[]", "Yes", "Store parsed skill list."),
            ("education", "EducationEntry[]", "Yes", "Map years to dates if backend keeps Date fields."),
            ("experience", "ExperienceSummary", "No", "Map entries into profile experience rows."),
            ("projects", "ProjectItem[]", "Yes", "Map `name -> title`, `technologies -> techStack`."),
            ("certifications", "string[]", "Yes", "Persist as-is."),
            ("education_level", "olevel|bachelor|master|phd", "No", "Normalize into backend enum."),
            ("years_experience", "number", "No", "Persist as denormalized summary for matching."),
            ("raw_text", "string", "No", "Strongly recommended to store in `Resume.parsedText`."),
            ("portfolio", "object", "No", "Schema supports it conceptually, but current parser does not populate it."),
        ],
    )

    add_heading(doc, "3.2 Match Candidates", 2)
    add_body(doc, "JSON endpoint for job-to-candidate ranking.", bold_label="AI route")
    add_code_block(doc, ["POST /match/", "Content-Type: application/json"])
    add_body(doc, "Single job payload plus candidate array.", bold_label="Request shape")
    add_code_block(
        doc,
        [
            "{",
            '  "job": {',
            '    "job_id": "<backend job id>",',
            '    "title": "Backend Engineer",',
            '    "description": "Full job description text",',
            '    "required_skills": ["node.js", "mongodb"],',
            '    "preferred_skills": ["redis"],',
            '    "required_experience_years": 3,',
            '    "required_education_level": "bachelor"',
            "  },",
            '  "candidates": [',
            '    {',
            '      "candidate_id": "<candidate profile id>",',
            '      "full_name": "Jane Doe",',
            '      "email": "jane@example.com",',
            '      "skills": ["node.js", "mongodb"],',
            '      "years_experience": 4,',
            '      "education_level": "bachelor",',
            '      "raw_text": "optional but recommended raw resume text"',
            "    }",
            "  ]",
            "}",
        ],
    )
    add_body(doc, "The ranking response fields are listed below.", bold_label="Response shape")
    add_four_col_table(
        doc,
        ["Field", "Type", "Meaning", "Persistence guidance"],
        [
            ("job_id", "string", "Echo of input job id", "Use to correlate the AI response."),
            ("ranked_candidates[].candidate_id", "string", "Echo of input candidate id", "Map back to `CandidateProfile._id`."),
            ("ranked_candidates[].total_score", "0..1 float", "Final weighted score", "Convert to percent or change DB schema."),
            ("ranked_candidates[].matched_skills", "string[]", "Required skills found", "Persist to `MatchResult.matchedSkills`."),
            ("ranked_candidates[].missing_skills", "string[]", "Required skills missing", "Persist to `MatchResult.missingSkills`."),
            ("ranked_candidates[].score_breakdown", "object", "skills/experience/semantic/education subscores", "Persist with clear naming."),
            ("ranked_candidates[].reasons", "string[]", "Human-readable reasons", "Good for recruiter UI detail."),
            ("ranked_candidates[].readable_summary", "string", "Generated narrative summary", "Persist to `MatchResult.explanation`."),
        ],
    )

    add_heading(doc, "4. Backend APIs to Build Around the AI Service", 1)
    add_body(doc, "The backend APIs below are the ones the frontend should call. They wrap the AI service and hide the AI-specific contracts.")
    add_three_col_table(
        doc,
        ["Backend Endpoint", "Purpose", "Backend action"],
        [
            ("POST /api/resumes", "Upload a resume", "Store file metadata, call AI `/parse`, persist parsed profile, return resume/profile ids and parse status."),
            ("GET /api/resumes/:resumeId", "Resume detail", "Return resume metadata plus parse status and any parse error."),
            ("GET /api/candidate-profiles/:profileId", "Candidate profile", "Return normalized AI-parsed profile combined with manual edits."),
            ("POST /api/jobs", "Create job requirement", "Persist normalized job requirement data in backend shape."),
            ("POST /api/jobs/:jobId/match", "Run a match", "Load job plus candidate pool, call AI `/match/`, upsert `MatchResult` rows."),
            ("GET /api/jobs/:jobId/matches", "List ranked matches", "Return persisted match results, optionally with filters and shortlist state."),
        ],
    )
    add_bullet(doc, "If resume parsing is expected to be slow, `POST /api/resumes` should be asynchronous: create the Resume row first, set `parseStatus=processing`, then complete parsing in a job worker.")
    add_bullet(doc, "If the team wants simpler first delivery, synchronous parsing is acceptable, but the backend should still update the parse status fields consistently.")

    add_heading(doc, "5. Required Field Mapping Layer", 1)
    add_body(doc, "One adapter module should own every transformation between backend models and AI payloads. Do not spread this mapping across controllers.")

    add_heading(doc, "5.1 JobRequirement -> AI JobInput", 2)
    add_four_col_table(
        doc,
        ["Backend field", "AI field", "Rule", "Important note"],
        [
            ("_id", "job_id", "Stringify ObjectId", "AI treats ids as opaque strings."),
            ("jobTitle", "title", "Copy directly", "Keep concise but do not omit description."),
            ("description", "description", "Copy directly", "This text powers semantic matching."),
            ("requiredSkills", "required_skills", "Lowercase/trim each value", "Avoid duplicates."),
            ("preferredSkills", "preferred_skills", "Lowercase/trim each value", "Avoid duplicates."),
            ("experienceYears", "required_experience_years", "Copy as number", "Defaults to 0."),
            ("educationLevel", "required_education_level", "Map enum", "`Any` should become `null`, not a string."),
        ],
    )

    add_heading(doc, "5.2 CandidateProfile + Resume -> AI CandidateInput", 2)
    add_four_col_table(
        doc,
        ["Backend field", "AI field", "Rule", "Important note"],
        [
            ("CandidateProfile._id", "candidate_id", "Stringify ObjectId", "Must stay stable across matching calls."),
            ("personalInfo.fullName", "full_name", "Copy directly", "Optional field."),
            ("personalInfo.email", "email", "Copy directly", "Optional field."),
            ("skills + manuallyAddedSkills", "skills", "Merge, trim, lowercase, de-duplicate", "Manual additions should not replace parsed skills."),
            ("Derived years total", "years_experience", "Persist a numeric summary", "AI expects a number, not experience entries."),
            ("Normalized education level", "education_level", "Map enum to lowercase literal", "Use `null` when unknown."),
            ("Resume.parsedText", "raw_text", "Send when available", "This is strongly recommended because AI matching uses richer embeddings when raw text is present."),
        ],
    )

    add_heading(doc, "5.3 AI ParsedCandidate -> CandidateProfile + Resume", 2)
    add_bullet(doc, "Store `raw_text` into `Resume.parsedText` exactly once per parsed resume version.")
    add_bullet(doc, "Write parsed contact fields into `CandidateProfile.personalInfo`.")
    add_bullet(doc, "Map `experience.entries[*].role -> jobTitle` and convert `start_year` / `end_year` into dates if the backend schema keeps Date types.")
    add_bullet(doc, "Map `projects[*].name -> title` and `technologies -> techStack`.")
    add_bullet(doc, "Persist `years_experience` and normalized `education_level` even if they are derivable later; they simplify matching payload creation.")

    add_heading(doc, "5.4 AI MatchResponse -> MatchResult", 2)
    add_four_col_table(
        doc,
        ["AI field", "Backend field", "Rule", "Mismatch to fix"],
        [
            ("total_score", "matchScore", "Either multiply by 100 or change schema to 0..1", "Current backend model says 0..100."),
            ("matched_skills", "matchedSkills", "Copy array", "Keep normalized casing strategy consistent."),
            ("missing_skills", "missingSkills", "Copy array", "Keep normalized casing strategy consistent."),
            ("readable_summary", "explanation", "Copy string", "Best field for recruiter-facing summary."),
            ("score_breakdown.skills_score", "scoreBreakdown.skills", "Rename key", "Current key names differ."),
            ("score_breakdown.experience_score", "scoreBreakdown.experience", "Rename key", "Current key names differ."),
            ("score_breakdown.education_score", "scoreBreakdown.education", "Rename key", "Current key names differ."),
            ("score_breakdown.semantic_score", "scoreBreakdown.semantic", "Rename key", "Current key names differ."),
        ],
    )

    add_heading(doc, "6. Integration Mismatches to Resolve Early", 1)
    add_bullet(doc, "Education enum mismatch: AI uses `olevel|bachelor|master|phd`; backend currently uses `Any|Bachelor|Master|PhD`. Add explicit mapping both ways.")
    add_bullet(doc, "Score range mismatch: AI returns `total_score` in the range 0 to 1, while `MatchResult.matchScore` currently expects 0 to 100.")
    add_bullet(doc, "Candidate raw text path: matching is materially better when `raw_text` is sent, but the current backend models do not clearly guarantee that path from `Resume.parsedText` into match requests.")
    add_bullet(doc, "Date granularity mismatch: AI parse outputs experience and education years, while the backend models currently use full Date fields.")
    add_bullet(doc, "Portfolio data is in the AI schema but is not currently populated by the parser; backend should treat it as optional and not depend on it.")

    add_heading(doc, "7. Suggested Backend Flow", 1)
    add_body(doc, "Resume upload flow", bold_label="Sequence")
    add_bullet(doc, "Create `Resume` row with `parseStatus=pending` and save file metadata.")
    add_bullet(doc, "Upload or store the binary file in your file store and keep the final URL.")
    add_bullet(doc, "Set `parseStatus=processing` and `parseStartedAt`.")
    add_bullet(doc, "Send the file to AI `POST /parse`.")
    add_bullet(doc, "Persist `Resume.parsedText` plus normalized `CandidateProfile` fields from the parse response.")
    add_bullet(doc, "Set `parseStatus=done` and `parseCompletedAt`, or `failed` with `parseError` if the AI call fails.")

    add_body(doc, "Matching flow", bold_label="Sequence")
    add_bullet(doc, "Load the target job and the candidate pool from backend storage.")
    add_bullet(doc, "Transform each backend record into AI `JobInput` and `CandidateInput` via one adapter module.")
    add_bullet(doc, "Call AI `POST /match/` with the normalized payload.")
    add_bullet(doc, "Upsert one `MatchResult` per candidate using unique key `{jobId, candidateId}`.")
    add_bullet(doc, "Return persisted ranked results to the frontend rather than proxying the raw AI response directly.")

    add_heading(doc, "8. Error Handling Rules", 1)
    add_bullet(doc, "If AI `/parse` returns HTTP 400, surface a clean backend validation error and keep the user-facing message understandable.")
    add_bullet(doc, "If AI `/parse` or `/match/` returns HTTP 500 or times out, mark the backend operation as failed or retryable; do not persist half-complete results as successful.")
    add_bullet(doc, "Log backend request ids alongside AI request payload identifiers such as job id and candidate ids to make debugging cross-service issues easier.")
    add_bullet(doc, "Do not swallow AI errors in the controller layer; convert them into explicit backend error objects.")

    add_heading(doc, "9. Recommended Internal Modules", 1)
    add_bullet(doc, "`services/aiClient` for outbound HTTP calls to the AI service.")
    add_bullet(doc, "`mappers/aiPayloadMapper` for JobRequirement and CandidateProfile to AI payload conversion.")
    add_bullet(doc, "`mappers/aiResponseMapper` for parse and match response persistence mapping.")
    add_bullet(doc, "`workers/resumeParsingWorker` if parsing runs asynchronously.")
    add_bullet(doc, "`services/matchService` to orchestrate candidate loading, AI call, and `MatchResult` upserts.")

    add_heading(doc, "10. Immediate Implementation Checklist", 1)
    add_bullet(doc, "Create enum mapping helpers for education levels in both directions.")
    add_bullet(doc, "Decide whether `MatchResult.matchScore` will store 0..1 or 0..100 and make the codebase consistent.")
    add_bullet(doc, "Ensure `Resume.parsedText` is saved and later forwarded as `raw_text` in match requests.")
    add_bullet(doc, "Add a derived `yearsExperience` value on the stored candidate profile if it is not already persisted.")
    add_bullet(doc, "Implement id-safe upserts for `{jobId, candidateId}` match rows.")
    add_bullet(doc, "Wrap all AI calls in one reusable client with base URL, timeout, and structured error handling.")

    add_heading(doc, "11. Final Recommendation", 1)
    add_body(
        doc,
        "The cleanest architecture is for the backend to expose stable product APIs and treat the current AI service as an internal scoring/parsing dependency. If the adapter layer is kept centralized, the team can change AI payloads later without forcing database or frontend rewrites."
    )

    doc.save(output_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    out_dir = root / "docs"
    out_dir.mkdir(parents=True, exist_ok=True)
    build_document(out_dir / "AI-Service-Backend-Handoff.docx")
